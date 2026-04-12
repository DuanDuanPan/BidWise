import os
import re
import struct
import time
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from docx_renderer.engine.figure_numbering import (
    FigureEntry,
    build_figure_registry,
    renumber_registry,
    replace_cross_references,
)
from docx_renderer.models.schemas import PageSetup, RenderResult, StyleMapping

# Fallback style names (built-in Word/python-docx English names)
_HEADING_FALLBACKS = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}

_HEADING_STYLE_KEYS = {
    1: "heading1",
    2: "heading2",
    3: "heading3",
    4: "heading4",
    5: "heading5",
    6: "heading6",
}

# Inline format pattern: bold+italic, bold, italic, inline code
# Order matters: *** must be matched before ** and *
_INLINE_PATTERN = re.compile(
    r"(\*\*\*|___)(.+?)\1"  # bold+italic
    r"|(\*\*|__)(.+?)\3"  # bold
    r"|(\*|_)(.+?)\5"  # italic
    r"|(`+)(.+?)\7"  # inline code
)

_IMAGE_PATTERN = re.compile(r"^!\[((?:[^\]\\]|\\.)*)\]\(([^)]+)\)\s*$")
_FENCED_CODE_START = re.compile(r"^(`{3,}|~{3,})(\w*)\s*$")

# Allowed image extensions
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}

# Monospace font for code fallback
_CODE_FONT = "Courier New"


class RendererError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code


def _style_exists(doc: Document, style_name: str) -> bool:
    """Check if a style name exists in the document."""
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


def _is_paragraph_style(doc: Document, style_name: str) -> bool:
    """Check if a style exists and is a PARAGRAPH style (not CHARACTER/TABLE/LIST)."""
    try:
        return doc.styles[style_name].type == WD_STYLE_TYPE.PARAGRAPH
    except KeyError:
        return False


def _resolve_paragraph_style(
    doc: Document,
    configured_name: Optional[str],
    fallback_name: str,
    warnings: list[str],
) -> Optional[str]:
    """Resolve a paragraph style: configured > fallback > None (with warning)."""
    if configured_name and _style_exists(doc, configured_name):
        if _is_paragraph_style(doc, configured_name):
            return configured_name
        warnings.append(
            f"样式 '{configured_name}' 不是段落样式 (类型不兼容)，使用 fallback '{fallback_name}'"
        )
    elif configured_name:
        warnings.append(f"样式 '{configured_name}' 在模板中不存在，使用 fallback '{fallback_name}'")
    if _is_paragraph_style(doc, fallback_name):
        return fallback_name
    if _style_exists(doc, fallback_name):
        warnings.append(f"Fallback 样式 '{fallback_name}' 不是段落样式，使用无样式直写")
        return None
    warnings.append(f"Fallback 样式 '{fallback_name}' 也不可用，使用无样式直写")
    return None


def _is_table_style(doc: Document, style_name: str) -> bool:
    """Check if a style exists and is a TABLE style."""
    try:
        return doc.styles[style_name].type == WD_STYLE_TYPE.TABLE
    except KeyError:
        return False


def _resolve_table_style(
    doc: Document,
    configured_name: Optional[str],
    warnings: list[str],
) -> Optional[str]:
    """Resolve a table style."""
    if configured_name and _style_exists(doc, configured_name):
        if _is_table_style(doc, configured_name):
            return configured_name
        warnings.append(
            f"表格样式 '{configured_name}' 不是表格样式 (类型不兼容)，使用无样式"
        )
        return None
    if configured_name:
        warnings.append(f"表格样式 '{configured_name}' 在模板中不存在")
    return None


def _get_style_key(style_mapping: Optional[StyleMapping], key: str) -> Optional[str]:
    """Get a style value from StyleMapping by key name."""
    if style_mapping is None:
        return None
    return getattr(style_mapping, key, None)


def _append_inline_runs(
    paragraph: Paragraph,
    text: str,
    style_mapping: Optional[StyleMapping],
    warnings: list[str],
    doc: Optional[Document] = None,
) -> None:
    """Parse inline Markdown formatting and append runs to a paragraph."""
    # Resolve inline code style once: only use if it is a CHARACTER style
    inline_code_style: Optional[str] = None
    code_style_name = _get_style_key(style_mapping, "code_block")
    if code_style_name and doc and _style_exists(doc, code_style_name):
        style_obj = doc.styles[code_style_name]
        if style_obj.type == WD_STYLE_TYPE.CHARACTER:
            inline_code_style = code_style_name

    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        # Add text before match as a plain run
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])

        if m.group(1):  # bold+italic (*** or ___)
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # bold (** or __)
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group(5):  # italic (* or _)
            run = paragraph.add_run(m.group(6))
            run.italic = True
        elif m.group(7):  # inline code (`)
            code_text = m.group(8)
            run = paragraph.add_run(code_text)
            if inline_code_style:
                run.style = doc.styles[inline_code_style]
            else:
                run.font.name = _CODE_FONT

        pos = m.end()

    # Add remaining text
    if pos < len(text):
        paragraph.add_run(text[pos:])
    elif pos == 0:
        # No inline formatting found, add as single run
        paragraph.add_run(text)


def _get_image_size(image_path: str) -> tuple[Optional[int], Optional[int]]:
    """Get image dimensions (width, height) in pixels using standard library only."""
    try:
        ext = Path(image_path).suffix.lower()
        with open(image_path, "rb") as f:
            if ext == ".png":
                # PNG: IHDR chunk at offset 16 contains width and height as 4-byte big-endian
                f.read(16)
                width_bytes = f.read(4)
                height_bytes = f.read(4)
                if len(width_bytes) == 4 and len(height_bytes) == 4:
                    width = struct.unpack(">I", width_bytes)[0]
                    height = struct.unpack(">I", height_bytes)[0]
                    return width, height
            elif ext in (".jpg", ".jpeg"):
                # JPEG: scan for SOF0/SOF2 marker
                f.read(2)  # skip SOI marker
                while True:
                    marker = f.read(2)
                    if len(marker) < 2:
                        break
                    if marker[0] != 0xFF:
                        break
                    if marker[1] in (0xC0, 0xC2):  # SOF0, SOF2
                        f.read(3)  # skip length and precision
                        h_bytes = f.read(2)
                        w_bytes = f.read(2)
                        if len(h_bytes) == 2 and len(w_bytes) == 2:
                            height = struct.unpack(">H", h_bytes)[0]
                            width = struct.unpack(">H", w_bytes)[0]
                            return width, height
                        break
                    else:
                        # Skip this segment
                        length_bytes = f.read(2)
                        if len(length_bytes) < 2:
                            break
                        length = struct.unpack(">H", length_bytes)[0]
                        f.read(length - 2)
    except Exception:
        pass
    return None, None


def _resolve_image_path(
    image_path_raw: str,
    project_path: Optional[str],
) -> Optional[str]:
    """Resolve and validate an image path, returning the resolved filesystem
    path on success or ``None`` when the image would be rejected.

    Validation gates (same order as ``_handle_image``):
    1. Extension whitelist
    2. Path resolution (relative / absolute)
    3. Security — must be under ``{project_path}/assets/``
    4. File existence
    """
    ext = Path(image_path_raw).suffix.lower()
    if ext not in _IMAGE_EXTENSIONS:
        return None

    if project_path and not os.path.isabs(image_path_raw):
        resolved = os.path.realpath(os.path.join(project_path, image_path_raw))
        assets_dir = os.path.realpath(os.path.join(project_path, "assets"))
        if not resolved.startswith(assets_dir + os.sep) and resolved != assets_dir:
            return None
    elif os.path.isabs(image_path_raw):
        if not project_path:
            return None
        resolved = os.path.realpath(image_path_raw)
        assets_dir = os.path.realpath(os.path.join(project_path, "assets"))
        if not resolved.startswith(assets_dir + os.sep) and resolved != assets_dir:
            return None
    else:
        return None

    if not os.path.exists(resolved):
        return None

    return resolved


def _handle_image(
    doc: Document,
    alt_text: str,
    image_path_raw: str,
    project_path: Optional[str],
    page_setup: Optional[PageSetup],
    warnings: list[str],
) -> bool:
    """Handle Markdown image syntax and insert into document."""
    resolved = _resolve_image_path(image_path_raw, project_path)

    if resolved is None:
        # Produce a human-readable warning (replicate the specific messages)
        ext = Path(image_path_raw).suffix.lower()
        if ext not in _IMAGE_EXTENSIONS:
            warnings.append(f"不支持的图片格式 '{ext}': {image_path_raw}")
        elif not project_path and not os.path.isabs(image_path_raw):
            warnings.append(f"无法解析图片路径 (缺少 projectPath): {image_path_raw}")
        elif not project_path:
            warnings.append(f"缺少 projectPath，无法校验绝对路径图片: {image_path_raw}")
        else:
            real = os.path.realpath(
                os.path.join(project_path, image_path_raw)
                if not os.path.isabs(image_path_raw)
                else image_path_raw
            )
            assets_dir = os.path.realpath(os.path.join(project_path, "assets"))
            if not real.startswith(assets_dir + os.sep) and real != assets_dir:
                if os.path.isabs(image_path_raw):
                    warnings.append(f"绝对路径图片不在 assets/ 目录下: {image_path_raw}")
                else:
                    warnings.append(f"图片路径不在 assets/ 目录下: {image_path_raw}")
            else:
                warnings.append(f"图片文件不存在: {real}")
        doc.add_paragraph(f"[图片未导出: {image_path_raw}]")
        return False

    # Insert image
    try:
        content_width_mm = 150.0
        if page_setup and page_setup.content_width_mm:
            content_width_mm = page_setup.content_width_mm

        content_width_emu = Mm(content_width_mm)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        shape = run.add_picture(resolved)

        # Scale down if the image's physical width exceeds content area
        if shape.width > content_width_emu:
            ratio = content_width_emu / shape.width
            shape.width = content_width_emu
            shape.height = int(shape.height * ratio)
    except Exception as e:
        warnings.append(f"图片插入失败: {image_path_raw}: {e}")
        doc.add_paragraph(f"[图片未导出: {image_path_raw}]")
        return False

    return True


def add_toc(
    doc: Document,
    title: str = "目录",
    toc_style: Optional[str] = None,
    warnings: Optional[list[str]] = None,
    first_heading_element=None,
) -> None:
    """Insert TOC field code before the first heading in the document."""
    if warnings is None:
        warnings = []

    # Add TOC title paragraph
    if title:
        if not toc_style:
            warnings.append("未配置 TOC 标题样式，使用 fallback 'Heading 1'")
        style = _resolve_paragraph_style(
            doc,
            toc_style,
            "Heading 1",
            warnings,
        )
        p = doc.add_paragraph(title, style=style)
    else:
        p = doc.add_paragraph()

    # Add TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)

    body = doc.element.body

    if first_heading_element is not None:
        # Insert TOC elements just before the first heading
        # Insert the field paragraph first (it will sit between title and heading)
        first_heading_element.addprevious(paragraph._p)
        # Then insert title before the field paragraph
        if title:
            paragraph._p.addprevious(p._p)
    else:
        # Fallback: insert at document beginning
        body.insert(0, paragraph._p)
        if title:
            body.insert(0, p._p)


def render_markdown_to_docx(
    markdown_content: str,
    output_path: str,
    template_path: Optional[str] = None,
    style_mapping: Optional[StyleMapping] = None,
    page_setup: Optional[PageSetup] = None,
    project_path: Optional[str] = None,
) -> RenderResult:
    start_time = time.perf_counter()
    warnings: list[str] = []

    if template_path:
        if not os.path.exists(template_path):
            raise RendererError("TEMPLATE_NOT_FOUND", f"Template not found: {template_path}")
        try:
            doc = Document(template_path)
        except (PackageNotFoundError, Exception) as e:
            raise RendererError(
                "DOCX_TEMPLATE_INVALID", f"Invalid docx template: {template_path}: {e}"
            ) from e
    else:
        doc = Document()

    # Build figure registry, prune images that will fail validation,
    # re-number survivors, then replace cross-references.
    lines = markdown_content.split("\n")
    figure_registry = build_figure_registry(lines)

    valid_figures: list[FigureEntry] = []
    for entry in figure_registry:
        img_match = _IMAGE_PATTERN.match(lines[entry.line_index])
        if img_match and _resolve_image_path(img_match.group(2), project_path) is not None:
            valid_figures.append(entry)
    figure_registry = renumber_registry(valid_figures)

    lines = replace_cross_references(lines, figure_registry, warnings)
    processed_content = "\n".join(lines)

    _parse_markdown(doc, processed_content, style_mapping, page_setup, project_path, warnings, figure_registry)

    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    doc.save(output_path)

    elapsed_ms = (time.perf_counter() - start_time) * 1000

    return RenderResult(
        output_path=output_path,
        render_time_ms=round(elapsed_ms, 2),
        warnings=warnings,
    )


def _add_code_block_paragraph(
    doc: Document,
    code_text: str,
    style_mapping: Optional[StyleMapping],
    warnings: list[str],
) -> None:
    """Add a fenced code block as a styled paragraph."""
    code_style = _get_style_key(style_mapping, "code_block")
    resolved_style = _resolve_paragraph_style(doc, code_style, "Normal", warnings) if code_style else None

    for line in code_text.split("\n"):
        p = doc.add_paragraph(style=resolved_style)
        run = p.add_run(line)
        if not resolved_style or resolved_style == "Normal":
            # Fallback: monospace + light gray shading
            run.font.name = _CODE_FONT
            run.font.size = Pt(9)
            # Add shading via XML
            rpr = run._r.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            rpr.append(shd)


def _add_figure_caption(
    doc: Document,
    figure_entry: FigureEntry,
    style_mapping: Optional[StyleMapping],
    warnings: list[str],
) -> None:
    """Insert a centered caption paragraph below the figure image."""
    configured = _get_style_key(style_mapping, "caption")
    resolved = _resolve_paragraph_style(doc, configured, "Caption", warnings)

    caption_text = f"{figure_entry.label}: {figure_entry.caption}"
    p = doc.add_paragraph(style=resolved)
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(caption_text)


def _parse_markdown(
    doc: Document,
    content: str,
    style_mapping: Optional[StyleMapping] = None,
    page_setup: Optional[PageSetup] = None,
    project_path: Optional[str] = None,
    warnings: Optional[list[str]] = None,
    figure_registry: Optional[list[FigureEntry]] = None,
) -> None:
    if warnings is None:
        warnings = []

    # Build lookup for figure entries by line index
    figure_by_line: dict[int, FigureEntry] = {}
    if figure_registry:
        for entry in figure_registry:
            figure_by_line[entry.line_index] = entry

    lines = content.split("\n")
    i = 0
    first_heading_element = None

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        fence_match = _FENCED_CODE_START.match(line)
        if fence_match:
            fence_char = fence_match.group(1)[0]
            fence_len = len(fence_match.group(1))
            code_lines: list[str] = []
            i += 1
            while i < len(lines):
                end_match = re.match(
                    rf"^ {{0,3}}{re.escape(fence_char)}{{{fence_len},}}\s*$",
                    lines[i],
                )
                if end_match:
                    i += 1
                    break
                code_lines.append(lines[i])
                i += 1
            _add_code_block_paragraph(doc, "\n".join(code_lines), style_mapping, warnings)
            continue

        # Image
        img_match = _IMAGE_PATTERN.match(line)
        if img_match:
            image_ok = _handle_image(
                doc,
                img_match.group(1),
                img_match.group(2),
                project_path,
                page_setup,
                warnings,
            )
            # Add figure caption only when the image was actually inserted
            figure_entry = figure_by_line.get(i)
            if figure_entry and image_ok:
                _add_figure_caption(
                    doc,
                    figure_entry,
                    style_mapping,
                    warnings,
                )
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            style_key = _HEADING_STYLE_KEYS.get(level, "heading1")
            configured = _get_style_key(style_mapping, style_key)
            fallback = _HEADING_FALLBACKS.get(level, "Heading 1")
            resolved = _resolve_paragraph_style(doc, configured, fallback, warnings)
            if resolved:
                p = doc.add_paragraph(style=resolved)
                _append_inline_runs(p, text, style_mapping, warnings, doc)
            else:
                p = doc.add_paragraph()
                _append_inline_runs(p, text, style_mapping, warnings, doc)
            if first_heading_element is None:
                first_heading_element = p._p
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(
            r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]
        ):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _parse_table(doc, table_lines, style_mapping, warnings)
            continue

        # Unordered list
        ul_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if ul_match:
            configured = _get_style_key(style_mapping, "list_bullet")
            resolved = _resolve_paragraph_style(doc, configured, "List Bullet", warnings)
            p = doc.add_paragraph(style=resolved)
            _append_inline_runs(p, ul_match.group(1), style_mapping, warnings, doc)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if ol_match:
            configured = _get_style_key(style_mapping, "list_number")
            resolved = _resolve_paragraph_style(doc, configured, "List Number", warnings)
            p = doc.add_paragraph(style=resolved)
            _append_inline_runs(p, ol_match.group(1), style_mapping, warnings, doc)
            i += 1
            continue

        # Empty line — skip
        if line.strip() == "":
            i += 1
            continue

        # Normal paragraph
        configured = _get_style_key(style_mapping, "body_text")
        resolved = _resolve_paragraph_style(doc, configured, "Normal", warnings) if configured else None
        p = doc.add_paragraph(style=resolved)
        _append_inline_runs(p, line, style_mapping, warnings, doc)
        i += 1

    # Insert TOC if there are headings
    if first_heading_element is not None:
        toc_style = _get_style_key(style_mapping, "toc")
        add_toc(
            doc,
            title="目录",
            toc_style=toc_style,
            warnings=warnings,
            first_heading_element=first_heading_element,
        )


def _parse_table(
    doc: Document,
    table_lines: list[str],
    style_mapping: Optional[StyleMapping] = None,
    warnings: Optional[list[str]] = None,
) -> None:
    if warnings is None:
        warnings = []

    rows = []
    for idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if idx == 1:
            # Skip separator row
            continue
        rows.append(cells)

    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)

    # Apply table style if configured
    table_style = _get_style_key(style_mapping, "table")
    resolved_table_style = _resolve_table_style(doc, table_style, warnings)
    if resolved_table_style:
        table.style = resolved_table_style

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                # Clear default paragraph and use inline formatting
                cell.text = ""
                p = cell.paragraphs[0]
                _append_inline_runs(p, cell_text, style_mapping, warnings if warnings else [], doc)
