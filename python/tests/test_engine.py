import os

import pytest
from docx import Document

from docx_renderer.engine.renderer import RendererError, render_markdown_to_docx


def test_render_heading_levels(tmp_path):
    output = str(tmp_path / "out.docx")
    md = "# H1\n## H2\n### H3"
    result = render_markdown_to_docx(md, output)
    assert os.path.exists(output)
    assert result.render_time_ms >= 0

    doc = Document(output)
    assert doc.paragraphs[0].text == "H1"
    assert doc.paragraphs[1].text == "H2"
    assert doc.paragraphs[2].text == "H3"


def test_render_mixed_content(tmp_path):
    output = str(tmp_path / "out.docx")
    md = "# Title\n\nSome text.\n\n- bullet\n\n1. numbered"
    result = render_markdown_to_docx(md, output)
    assert result.output_path == output

    doc = Document(output)
    texts = [p.text for p in doc.paragraphs]
    assert "Title" in texts
    assert "Some text." in texts
    assert "bullet" in texts
    assert "numbered" in texts


def test_render_table(tmp_path):
    output = str(tmp_path / "out.docx")
    md = "| Col1 | Col2 |\n| --- | --- |\n| A | B |\n| C | D |"
    render_markdown_to_docx(md, output)

    doc = Document(output)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Col1"
    assert table.cell(2, 1).text == "D"


def test_render_template_not_found(tmp_path):
    output = str(tmp_path / "out.docx")
    with pytest.raises(RendererError) as exc_info:
        render_markdown_to_docx("# Test", output, template_path="/no/such/file.docx")
    assert exc_info.value.code == "TEMPLATE_NOT_FOUND"


def test_render_template_invalid(tmp_path):
    output = str(tmp_path / "out.docx")
    bad_template = str(tmp_path / "bad.docx")
    with open(bad_template, "w") as f:
        f.write("not a docx")
    with pytest.raises(RendererError) as exc_info:
        render_markdown_to_docx("# Test", output, template_path=bad_template)
    assert exc_info.value.code == "DOCX_TEMPLATE_INVALID"


def test_render_with_valid_template(tmp_path):
    template = str(tmp_path / "tmpl.docx")
    Document().save(template)

    output = str(tmp_path / "out.docx")
    result = render_markdown_to_docx("# Templated", output, template_path=template)
    assert os.path.exists(result.output_path)


def test_render_creates_parent_dirs(tmp_path):
    output = str(tmp_path / "deep" / "nested" / "out.docx")
    render_markdown_to_docx("# Test", output)
    assert os.path.exists(output)


def test_render_empty_content(tmp_path):
    output = str(tmp_path / "out.docx")
    result = render_markdown_to_docx("", output)
    assert os.path.exists(result.output_path)
