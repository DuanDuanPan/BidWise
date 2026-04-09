import os

from docx import Document
from fastapi.testclient import TestClient


def test_render_basic_heading(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Hello World",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["outputPath"] == tmp_output
    assert "renderTimeMs" in body["data"]
    assert os.path.exists(tmp_output)

    doc = Document(tmp_output)
    # TOC title + TOC field + heading = 3 paragraphs
    assert len(doc.paragraphs) == 3
    assert doc.paragraphs[0].text == "目录"  # TOC title
    assert doc.paragraphs[2].text == "Hello World"  # The heading


def test_render_paragraph(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "This is a paragraph.",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    assert doc.paragraphs[0].text == "This is a paragraph."


def test_render_unordered_list(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "- item one\n- item two",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200

    doc = Document(tmp_output)
    assert doc.paragraphs[0].text == "item one"
    assert doc.paragraphs[0].style.name == "List Bullet"


def test_render_ordered_list(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "1. first\n2. second",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200

    doc = Document(tmp_output)
    assert doc.paragraphs[0].text == "first"
    assert doc.paragraphs[0].style.name == "List Number"


def test_render_table(client: TestClient, tmp_output: str):
    md = "| A | B |\n| --- | --- |\n| 1 | 2 |"
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": md,
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200

    doc = Document(tmp_output)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.cell(0, 0).text == "A"
    assert table.cell(0, 1).text == "B"
    assert table.cell(1, 0).text == "1"
    assert table.cell(1, 1).text == "2"


def test_render_template_not_found(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": tmp_output,
            "templatePath": "/nonexistent/template.docx",
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is False
    assert body["error"]["code"] == "TEMPLATE_NOT_FOUND"


def test_render_template_invalid(client: TestClient, tmp_output: str, tmp_path):
    invalid_template = str(tmp_path / "invalid.docx")
    with open(invalid_template, "w") as f:
        f.write("not a docx file")

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": tmp_output,
            "templatePath": invalid_template,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is False
    assert body["error"]["code"] == "DOCX_TEMPLATE_INVALID"


def test_render_empty_content(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert os.path.exists(tmp_output)


def test_render_with_valid_template(client: TestClient, tmp_output: str, tmp_path):
    template_path = str(tmp_path / "template.docx")
    doc = Document()
    doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Templated",
            "outputPath": tmp_output,
            "templatePath": template_path,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True


def test_render_camel_case_response(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    body = response.json()
    data = body["data"]
    assert "outputPath" in data
    assert "renderTimeMs" in data
    assert "warnings" in data
    assert "output_path" not in data
    assert "render_time_ms" not in data


def test_render_accepts_camel_case_style_mapping(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "styleMapping": {"heading1": "标题 1", "bodyText": "正文"},
            "pageSetup": {"contentWidthMm": 150},
            "projectPath": "/tmp/project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert isinstance(body["data"]["warnings"], list)


# === Task 3: Style mapping & warnings tests ===


def test_style_mapping_missing_style_generates_warning(client: TestClient, tmp_output: str):
    """When configured style doesn't exist, warning is recorded and fallback used."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Heading",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "styleMapping": {"heading1": "不存在的样式"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    warnings = body["data"]["warnings"]
    # Should contain warning about missing style
    assert any("不存在的样式" in w for w in warnings)

    doc = Document(tmp_output)
    # Heading should still be rendered (using fallback)
    heading_texts = [p.text for p in doc.paragraphs if p.text == "Heading"]
    assert len(heading_texts) == 1


def test_no_template_uses_fallback_styles(client: TestClient, tmp_output: str):
    """Without template, built-in styles are used as fallback."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# H1\n\n## H2\n\nParagraph text",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    # Should have headings with built-in styles
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles or any("Heading" in s for s in styles)


def test_style_mapping_with_valid_template(client: TestClient, tmp_output: str, tmp_path):
    """Custom styles from template are used when they exist."""
    # Create a template with a custom style
    template_path = str(tmp_path / "custom-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("MyHeading", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Custom Styled",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"heading1": "MyHeading"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    # Find the heading paragraph (skip TOC)
    heading_paras = [p for p in doc.paragraphs if p.text == "Custom Styled"]
    assert len(heading_paras) == 1
    assert heading_paras[0].style.name == "MyHeading"


# === Task 4: Inline formatting & code block tests ===


def test_inline_bold(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "This is **bold** text",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    runs = para.runs
    assert any(r.bold for r in runs)
    assert para.text == "This is bold text"


def test_inline_italic(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "This is *italic* text",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    assert any(r.italic for r in para.runs)
    assert para.text == "This is italic text"


def test_inline_code(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "Use `print()` function",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    code_runs = [r for r in para.runs if r.font.name == "Courier New"]
    assert len(code_runs) == 1
    assert code_runs[0].text == "print()"


def test_inline_bold_italic(client: TestClient, tmp_output: str):
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "This is ***bold italic*** text",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    bold_italic_runs = [r for r in para.runs if r.bold and r.italic]
    assert len(bold_italic_runs) == 1
    assert bold_italic_runs[0].text == "bold italic"


def test_fenced_code_block(client: TestClient, tmp_output: str):
    md = "```python\ndef hello():\n    print('hi')\n```"
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": md,
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    # Code block should have monospace font runs
    code_texts = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name == "Courier New":
                code_texts.append(r.text)
    assert "def hello():" in code_texts
    assert "    print('hi')" in code_texts


def test_empty_fenced_code_block(client: TestClient, tmp_output: str):
    md = "```\n```"
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": md,
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True


def test_fenced_code_block_with_tilde(client: TestClient, tmp_output: str):
    md = "~~~\ncode here\n~~~"
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": md,
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    found_code = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text == "code here" and r.font.name == "Courier New":
                found_code = True
    assert found_code


# === Task 5: Image insertion & asset boundary tests ===


def test_image_insertion_success(client: TestClient, tmp_output: str, tmp_path):
    """Image under project assets/ should be inserted successfully."""
    # Setup: create project structure with assets/
    project_path = str(tmp_path / "project")
    assets_dir = os.path.join(project_path, "assets")
    os.makedirs(assets_dir, exist_ok=True)
    # Copy test image
    import shutil

    fixture_image = os.path.join(os.path.dirname(__file__), "fixtures", "images", "test-image.png")
    shutil.copy(fixture_image, os.path.join(assets_dir, "diagram.png"))

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![diagram](assets/diagram.png)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    # No image-related warnings
    image_warnings = [w for w in body["data"]["warnings"] if "图片" in w]
    assert len(image_warnings) == 0


def test_image_missing_generates_warning(client: TestClient, tmp_output: str, tmp_path):
    """Missing image file generates warning and placeholder text."""
    project_path = str(tmp_path / "project")
    os.makedirs(os.path.join(project_path, "assets"), exist_ok=True)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![missing](assets/missing.png)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    placeholder_texts = [p.text for p in doc.paragraphs if "图片未导出" in p.text]
    assert len(placeholder_texts) >= 1
    assert any("图片不存在" in w or "missing" in w for w in body["data"]["warnings"])


def test_image_path_traversal_rejected(client: TestClient, tmp_output: str, tmp_path):
    """Path traversal attempts are blocked."""
    project_path = str(tmp_path / "project")
    os.makedirs(os.path.join(project_path, "assets"), exist_ok=True)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![evil](../../../etc/passwd)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    doc = Document(tmp_output)
    assert any("图片未导出" in p.text for p in doc.paragraphs)


def test_image_extension_whitelist(client: TestClient, tmp_output: str, tmp_path):
    """Non-whitelisted extensions are rejected."""
    project_path = str(tmp_path / "project")
    os.makedirs(os.path.join(project_path, "assets"), exist_ok=True)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![svg](assets/diagram.svg)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    warnings = body["data"]["warnings"]
    assert any("不支持的图片格式" in w for w in warnings)


def test_image_width_scaling(client: TestClient, tmp_output: str, tmp_path):
    """Large images should be scaled to content width."""
    project_path = str(tmp_path / "project")
    assets_dir = os.path.join(project_path, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    # Create a wide test image (2000x100 PNG)
    import struct
    import zlib

    width, height = 2000, 100
    raw = b""
    for _y in range(height):
        raw += b"\x00"
        for _x in range(width):
            raw += b"\xff\x00\x00\xff"
    compressed = zlib.compress(raw)

    def chunk(chunk_type, data):
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    png = b"\x89PNG\r\n\x1a\n"
    png += chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0))
    png += chunk(b"IDAT", compressed)
    png += chunk(b"IEND", b"")

    with open(os.path.join(assets_dir, "wide.png"), "wb") as f:
        f.write(png)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![wide](assets/wide.png)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
            "pageSetup": {"contentWidthMm": 150},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert os.path.exists(tmp_output)


# === Task 6: TOC field code tests ===


def test_toc_generated_when_headings_present(client: TestClient, tmp_output: str):
    """TOC should be auto-generated when document has headings."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Chapter 1\n\nText\n\n## Section 1.1",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    # First paragraph should be TOC title "目录"
    assert doc.paragraphs[0].text == "目录"


def test_toc_not_generated_without_headings(client: TestClient, tmp_output: str):
    """TOC should NOT be generated when no headings present."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "Just plain text\n\nMore text",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    assert doc.paragraphs[0].text == "Just plain text"
    assert not any(p.text == "目录" for p in doc.paragraphs)


def test_toc_xml_structure(client: TestClient, tmp_output: str):
    """Verify TOC field code XML structure."""
    from docx.oxml.ns import qn

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Heading One\n\n## Heading Two",
            "outputPath": tmp_output,
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)

    # Find the TOC field paragraph (second paragraph after "目录" title)
    toc_field_para = doc.paragraphs[1]
    runs = toc_field_para._p.findall(qn("w:r"))

    # Should have fldChar elements for begin, separate, and end
    fld_chars = []
    instr_texts = []
    for run in runs:
        for fc in run.findall(qn("w:fldChar")):
            fld_chars.append(fc.get(qn("w:fldCharType")))
        for it in run.findall(qn("w:instrText")):
            instr_texts.append(it.text)

    assert "begin" in fld_chars
    assert "separate" in fld_chars
    assert "end" in fld_chars
    assert any("TOC" in t for t in instr_texts)


def test_toc_style_fallback_warning(client: TestClient, tmp_output: str):
    """When TOC style doesn't exist, should fallback with warning."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "styleMapping": {"toc": "NonExistentTOCStyle"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    warnings = body["data"]["warnings"]
    assert any("TOC" in w or "toc" in w.lower() for w in warnings)


def test_inline_code_uses_style_mapping(client: TestClient, tmp_output: str, tmp_path):
    """Inline code should use codeBlock style from style_mapping when available in template."""
    # Create template with a character style named "CodeChar"
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "code-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("CodeChar", WD_STYLE_TYPE.CHARACTER)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "Use `print()` here",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"codeBlock": "CodeChar"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    code_runs = [r for r in para.runs if r.text == "print()"]
    assert len(code_runs) == 1
    assert code_runs[0].style.name == "CodeChar"


def test_inline_code_fallback_without_style(client: TestClient, tmp_output: str):
    """Inline code should fallback to Courier New when codeBlock style does not exist."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "Use `print()` here",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "styleMapping": {"codeBlock": "NonExistentCodeStyle"},
        },
    )
    assert response.status_code == 200
    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    code_runs = [r for r in para.runs if r.font.name == "Courier New"]
    assert len(code_runs) == 1
    assert code_runs[0].text == "print()"


def test_image_symlink_outside_assets_rejected(client: TestClient, tmp_output: str, tmp_path):
    """Symlink inside assets/ pointing outside project should be rejected."""
    project_path = str(tmp_path / "project")
    assets_dir = os.path.join(project_path, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    # Create a file outside the project
    outside_file = str(tmp_path / "secret.png")
    with open(outside_file, "wb") as f:
        # Minimal valid PNG header
        f.write(b"\x89PNG\r\n\x1a\n")

    # Create symlink inside assets pointing to outside file
    symlink_path = os.path.join(assets_dir, "linked.png")
    os.symlink(outside_file, symlink_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![linked](assets/linked.png)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    # Should have a warning about the image being outside assets
    warnings = body["data"]["warnings"]
    assert any("assets" in w for w in warnings)
    # Should have placeholder text, not the image content
    doc = Document(tmp_output)
    assert any("图片未导出" in p.text for p in doc.paragraphs)


def test_warnings_not_blocking_render(client: TestClient, tmp_output: str):
    """Warnings should not block rendering — doc should still be produced."""
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Title\n\nSome text\n- list item",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "styleMapping": {
                "heading1": "NonExistent1",
                "bodyText": "NonExistent2",
                "listBullet": "NonExistent3",
            },
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert os.path.exists(tmp_output)
    # Warnings should be recorded
    assert len(body["data"]["warnings"]) > 0


# === Regression tests for blocking fixes ===


def test_inline_code_paragraph_style_falls_back_to_monospace(
    client: TestClient, tmp_output: str, tmp_path
):
    """When codeBlock maps to a PARAGRAPH style, inline code must fall back to monospace font."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "para-code-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("ParaCode", WD_STYLE_TYPE.PARAGRAPH)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "Use `print()` here",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"codeBlock": "ParaCode"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    para = doc.paragraphs[0]
    code_runs = [r for r in para.runs if r.text == "print()"]
    assert len(code_runs) == 1
    # Should fall back to Courier New, not crash with ValueError
    assert code_runs[0].font.name == "Courier New"


def test_toc_inserted_before_first_heading_not_at_start(
    client: TestClient, tmp_output: str, tmp_path
):
    """TOC should be inserted before the first heading, preserving any template preamble."""
    # Create a template with preamble paragraphs (simulating a cover page)
    template_path = str(tmp_path / "preamble-template.docx")
    template_doc = Document()
    template_doc.add_paragraph("公司名称")
    template_doc.add_paragraph("封面页")
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# 第一章\n\n正文内容",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    texts = [p.text for p in doc.paragraphs]
    # Template preamble should come first
    assert texts[0] == "公司名称"
    assert texts[1] == "封面页"
    # Then TOC title
    assert texts[2] == "目录"
    # Then TOC field (empty text), then heading
    heading_idx = texts.index("第一章")
    assert heading_idx > 2


def test_image_width_scaled_to_content_width(client: TestClient, tmp_output: str, tmp_path):
    """Large images must be scaled down so their width equals content_width_mm."""
    import struct
    import zlib

    from docx.shared import Mm

    project_path = str(tmp_path / "project")
    assets_dir = os.path.join(project_path, "assets")
    os.makedirs(assets_dir, exist_ok=True)

    # Create a wide 2000x100 PNG (at default 96 DPI, ~528mm wide, well over 150mm)
    width, height = 2000, 100
    raw = b""
    for _y in range(height):
        raw += b"\x00"
        for _x in range(width):
            raw += b"\xff\x00\x00\xff"
    compressed = zlib.compress(raw)

    def chunk(chunk_type, data):
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    png = b"\x89PNG\r\n\x1a\n"
    png += chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0))
    png += chunk(b"IDAT", compressed)
    png += chunk(b"IEND", b"")

    with open(os.path.join(assets_dir, "wide.png"), "wb") as f:
        f.write(png)

    content_width = 150
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "![wide](assets/wide.png)",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "projectPath": project_path,
            "pageSetup": {"contentWidthMm": content_width},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True

    doc = Document(tmp_output)
    # Find the inline shape in the document
    from docx.oxml.ns import qn

    inline_shapes = doc.inline_shapes
    assert len(inline_shapes) == 1
    shape = inline_shapes[0]
    expected_width = Mm(content_width)
    # Allow 1% tolerance for rounding
    assert abs(shape.width - expected_width) / expected_width < 0.01


# === Regression: CHARACTER style mapped to paragraph-level key ===


def test_character_style_mapped_to_heading_falls_back(
    client: TestClient, tmp_output: str, tmp_path
):
    """When heading1 maps to a CHARACTER style, must fall back instead of raising ValueError."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "char-heading-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("CharStyle", WD_STYLE_TYPE.CHARACTER)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Heading with char style\n\nBody text here.",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"heading1": "CharStyle", "bodyText": "CharStyle"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    # Should have warnings about incompatible style type
    warnings = body["data"]["warnings"]
    assert any("不是段落样式" in w for w in warnings)
    # Document should still be generated with fallback styles
    assert os.path.exists(tmp_output)
    doc = Document(tmp_output)
    heading_texts = [p.text for p in doc.paragraphs if p.text == "Heading with char style"]
    assert len(heading_texts) == 1


def test_character_style_mapped_to_list_falls_back(
    client: TestClient, tmp_output: str, tmp_path
):
    """When listBullet maps to a CHARACTER style, must fall back."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "char-list-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("ListChar", WD_STYLE_TYPE.CHARACTER)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "- bullet one\n- bullet two",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"listBullet": "ListChar"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    warnings = body["data"]["warnings"]
    assert any("不是段落样式" in w for w in warnings)


# === Regression: absolute image path bypass when project_path is None ===


def test_absolute_image_path_rejected_without_project_path(
    client: TestClient, tmp_output: str, tmp_path
):
    """Absolute image path must be rejected when projectPath is omitted."""
    # Create a real PNG file at an absolute path
    secret_image = str(tmp_path / "secret.png")
    with open(secret_image, "wb") as f:
        f.write(b"\x89PNG\r\n\x1a\n")

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": f"![stolen]({secret_image})",
            "outputPath": tmp_output,
            "projectId": "test-project",
            # projectPath deliberately omitted
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    # Should have a warning about missing projectPath
    warnings = body["data"]["warnings"]
    assert any("projectPath" in w or "缺少" in w for w in warnings)
    # Should NOT have embedded the image — placeholder only
    doc = Document(tmp_output)
    assert any("图片未导出" in p.text for p in doc.paragraphs)


# === Regression: route handler catch-all for unexpected exceptions ===


def test_toc_character_style_falls_back(client: TestClient, tmp_output: str, tmp_path):
    """When toc maps to a CHARACTER style, must fall back to Heading 1 instead of raising ValueError."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "char-toc-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("TocChar", WD_STYLE_TYPE.CHARACTER)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Heading\n\nBody text",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"toc": "TocChar"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    warnings = body["data"]["warnings"]
    assert any("不是段落样式" in w for w in warnings)
    doc = Document(tmp_output)
    assert any(p.text == "目录" for p in doc.paragraphs)


def test_table_character_style_falls_back(client: TestClient, tmp_output: str, tmp_path):
    """When table maps to a CHARACTER style, must fall back instead of raising ValueError."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "char-table-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("TableChar", WD_STYLE_TYPE.CHARACTER)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "| Col1 | Col2 |\n| --- | --- |\n| A | B |",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"table": "TableChar"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    warnings = body["data"]["warnings"]
    assert any("不是表格样式" in w for w in warnings)
    doc = Document(tmp_output)
    assert len(doc.tables) == 1


def test_table_paragraph_style_falls_back(client: TestClient, tmp_output: str, tmp_path):
    """When table maps to a PARAGRAPH style, must fall back instead of raising ValueError."""
    from docx.enum.style import WD_STYLE_TYPE

    template_path = str(tmp_path / "para-table-template.docx")
    template_doc = Document()
    template_doc.styles.add_style("TablePara", WD_STYLE_TYPE.PARAGRAPH)
    template_doc.save(template_path)

    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "| Col1 | Col2 |\n| --- | --- |\n| A | B |",
            "outputPath": tmp_output,
            "projectId": "test-project",
            "templatePath": template_path,
            "styleMapping": {"table": "TablePara"},
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    warnings = body["data"]["warnings"]
    assert any("不是表格样式" in w for w in warnings)
    doc = Document(tmp_output)
    assert len(doc.tables) == 1


def test_unexpected_error_returns_structured_response(client: TestClient, tmp_output: str):
    """Unexpected exceptions must return structured error, not raw 500."""
    # Trigger an error by providing an output path that can't be created
    # (on most systems, /dev/null/foo is invalid)
    response = client.post(
        "/api/render-documents",
        json={
            "markdownContent": "# Test",
            "outputPath": "/dev/null/impossible/path/out.docx",
            "projectId": "test-project",
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is False
    assert body["error"]["code"] in ("RENDER_UNEXPECTED",)
    assert len(body["error"]["message"]) > 0
